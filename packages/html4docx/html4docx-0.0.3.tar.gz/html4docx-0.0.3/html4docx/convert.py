from docx import Document
from bs4 import BeautifulSoup
import re
import subprocess
import os
import base64

# from PIL import Image

# import io
import mammoth


# To convert emf type images to png (html does not support emf/wmf images)
# To install inkscape in linux
# sudo apt update
# sudo apt install inkscape
def convert_emf_to_png(emf_path, png_path):
    subprocess.run(
        [
            "inkscape",
            "-e",
            png_path,
            emf_path,
        ]
    )


# Callback to retrieve img actual width & height shown in docx file
def dim_callback(match):
    # CX
    cx_pat = 'cx="[0-9]+"'
    cx_match = re.search(cx_pat, match).group()
    cx_ = cx_match.split("=")
    cx = cx_[len(cx_) - 1].replace('"', "")
    # CY
    cy_pat = 'cy="[0-9]+"'
    cy_match = re.search(cy_pat, match).group()
    cy_ = cy_match.split("=")
    cy = cy_[len(cy_) - 1].replace('"', "")
    return int(cx), int(cy)


# Method to extract all img-dimensions in docx file
def extract_image_dimensions(docx_path):
    doc = Document(docx_path)
    paragraphs = []
    for idx, par in enumerate(doc.paragraphs):
        if par._p.xml.__contains__('w:val="Heading1"'):
            paragraphs = doc.paragraphs[idx : len(doc.paragraphs)]
            break
    content = " ".join([str(par._p.xml) for par in paragraphs])
    img_dims_ls = re.findall(
        """<wp:extent[\s]*cx="[0-9]*"[\s]*cy="[0-9]*"[\s]*\/>""", content
    )
    img_dims = []
    for img_dim in img_dims_ls:
        cx, cy = dim_callback(img_dim)
        img_dims.append((cx, cy))
    return img_dims


# Method to extract all imgs in docx file
def img_type_conversion(html, filename, is_convert=False):
    inch_const = 914400
    px_const = 96
    img_dims = extract_image_dimensions(filename)
    soup = BeautifulSoup(html, "html.parser")
    img_tags = soup.find_all("img")
    img_tags = list(reversed(img_tags))
    img_dims = list(reversed(img_dims))
    for img_id, img_tag in enumerate(img_tags):
        if len(img_dims) > img_id:
            cx, cy = img_dims[img_id]
            # img_tag["cx"] = cx
            # img_tag["cy"] = cy
            img_tag["width"] = f"{round((cx / inch_const) * px_const, 2)}px"
            img_tag["height"] = f"{round((cy / inch_const) * px_const, 2)}px"
        img_src = img_tag.get("src")
        img_type = img_src.split(";")[0]
        if img_type.__contains__("x-emf") and is_convert:
            img_data = img_src.split("base64,")[1]
            img_dec_data = base64.b64decode(img_data)
            emf_path = f"img_conversions/Figure_{img_id}.emf"
            png_path = f"img_conversions"
            png_img = ""
            # emf_image = io.BytesIO(img_dec_data)
            # emf_img = Image.open(emf_image)
            # actual_width = emf_img.width
            # actual_height = emf_img.height
            with open(emf_path, "wb") as img_file:
                img_file.write(img_dec_data)
            convert_emf_to_png(emf_path, png_path)
            if os.path.isfile(f"img_conversions/Figure_{img_id}.png"):
                with open(f"img_conversions/Figure_{img_id}.png", "rb") as png_file:
                    png_img = base64.b64encode(png_file.read())

                if len(png_img) > 0:
                    print(f"Converted Img {img_id}")
                    img_tag["src"] = f"data:image/png;base64,{png_img.decode()}"
                    # img_tag["style"] = f"width:{actual_width};height:{actual_height};"
    return str(soup)


# Method to convert docx to html
def docx_to_html(filename, is_convert=False):
    try:
        dir_path = "img_conversions"
        if not os.path.isdir(dir_path):
            os.mkdir(dir_path)
        with open(filename, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)

        conv_html = img_type_conversion(result.value, filename, is_convert)
        return conv_html
    except Exception as exp:
        return f"Error occured as {exp}"
