import os
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from PIL import Image

def create_word_document(folder_path, output_docx):
    document = Document()

    # Get a list of all image files in the specified folder
    image_files = [f for f in os.listdir(folder_path) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]

    for image_file in image_files:
        # Add a paragraph with the image name
        document.add_paragraph(image_file)

        # Add the image to the document
        image_path = os.path.join(folder_path, image_file)
        document.add_picture(image_path, width=Inches(4))  # Adjust width as needed

        # Add a newline between images
        document.add_paragraph()

    # Save the document
    document.save(output_docx)
    print("Successfully to Save the document")

def images_to_pdf(image_paths, pdf_path, title, copyright_notice):
    if not image_paths:
        print("No image files found.")
        return
    # Find the maximum width and height among all images
    max_width, max_height = max(Image.open(image_path).size for image_path in image_paths)
    c = canvas.Canvas(pdf_path, pagesize=(max_width, max_height))
    for idx, image_path in enumerate(image_paths):
        img = Image.open(image_path)
        img_width, img_height = img.size
        # Calculate the scaling factor to fit the image into the PDF
        scale_factor = min(max_width / img_width, max_height / img_height)
        # Calculate the position to center the image on the page
        x_position = (max_width - img_width * scale_factor) / 2
        y_position = (max_height - img_height * scale_factor) / 2
        # Draw the scaled image on the PDF
        c.drawInlineImage(image_path, x_position, y_position, img_width * scale_factor, img_height * scale_factor)
        # Add title in the center
        c.setFont("Helvetica", 16)
        c.drawCentredString(max_width / 2, max_height - 30, title)
        # Add copyright notice in the left corner
        c.setFont("Helvetica", 8)
        c.drawString(10, 10, copyright_notice)
        c.showPage()
    c.save()
 
