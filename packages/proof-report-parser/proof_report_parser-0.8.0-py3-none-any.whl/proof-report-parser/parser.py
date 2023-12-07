import datetime
import re
import traceback

from docx import Document
from docx.table import Table

from src.proof_report_parser.exceptions import RowNotFoundException, TableNotFoundException
from src.proof_report_parser.models import ReportData, Contact, CheckMan

datePattern = re.compile(r"(\d+).*?(\d+).*?(\d+).*?(\d+).*?(\d+).*?(\d+)")


def delSpace(mystring):
    if mystring is None:
        return ""
    else:
        return re.sub(r"\s+", "", mystring).rstrip()


def count_checked_unit(document: Document):
    _table_i = 0
    _count = 0
    while _table_i < len(document.tables):
        _cur_table = document.tables[_table_i]
        if _cur_table.rows[0].cells[0].text == "编号" and _cur_table.rows[0].cells[1].text == "建（构）筑物单体名称":
            _row_i = 1
            while _row_i < len(_cur_table.rows):
                if delSpace(_cur_table.rows[_row_i].cells[1].text) != "/":
                    _count += 1
                _row_i += 1
        _table_i += 1
    return _count


def count_check_point(document: Document) -> int:
    _count = 0
    _table_i = 0
    while _table_i < len(document.tables):  # 遍历table
        _cur_table = document.tables[_table_i]
        if _cur_table.rows[0].cells[0].text == "项目名称":  # 查找table标识
            _col_i = 0
            # _cur_row = _cur_table.rows[_table_i]
            while _col_i < len(_cur_table.rows[1].cells):  # 查找col标识:测试阻值并count
                _row_i = 0
                if "测试阻值" in _cur_table.rows[1].cells[_col_i].text:
                    while _row_i < len(_cur_table.rows):
                        try:
                            float(_cur_table.rows[_row_i].cells[_col_i].text)
                            _count += 1
                        except Exception:
                            pass
                        _row_i += 1
                if "接地电阻" in _cur_table.rows[1].cells[_col_i].text:
                    while _row_i < len(_cur_table.rows):
                        try:
                            float(_cur_table.rows[_row_i].cells[_col_i].text)
                            _count += 1
                        except Exception:
                            pass
                        _row_i += 1
                _col_i += 1
        _table_i = _table_i + 1
    return _count


def find_row_by_first_col_text(table: Table, text: str):
    _row_i = 0
    while _row_i < len(table.rows):
        if table.rows[_row_i].cells[0].text == text:
            return _row_i
        _row_i += 1
    raise RowNotFoundException(text)


def parseProofReport(document: Document) -> ReportData:
    reportData = ReportData()
    _mainTable = None
    i = 0
    while i < len(document.tables):
        if document.tables[i].rows[0].cells[0].text == "受检单位":
            _mainTable = document.tables[i]
            break
        i += 1
    if _mainTable is None:
        raise TableNotFoundException("MAIN TABLE")
    # 报告编号
    reportIdValue = document.tables[0].rows[0].cells[1].text
    reportData.reportFullId = delSpace(reportIdValue)
    # 受检单位
    reportData.checkedCompanyName = delSpace(_mainTable.rows[0].cells[1].text)
    # 项目名称
    reportData.projName = delSpace(_mainTable.rows[0].cells[5].text)
    # 检测类别
    i = 0
    target_text = "检测类别"
    while i < len(document.paragraphs):
        paragraph = document.paragraphs[i]
        if target_text in paragraph.text:
            reportData.checkType = delSpace(paragraph.text.replace(target_text, ""))
            break
        i += 1
    # 联系人
    _contact_row_index = find_row_by_first_col_text(_mainTable, "联系人")
    reportData.projContact = Contact(name=_mainTable.rows[_contact_row_index].cells[1].text,
                                     tel=_mainTable.rows[_contact_row_index].cells[6].text)
    # 防雷类别
    _proof_level_index = find_row_by_first_col_text(_mainTable, "防雷类别")
    reportData.proofLevel = delSpace(_mainTable.rows[_proof_level_index].cells[1].text)
    # 报告有效期
    _valid_date_index = find_row_by_first_col_text(_mainTable, "报告有效期")
    dateString = delSpace(_mainTable.rows[_valid_date_index].cells[6].text)
    match = datePattern.match(dateString)
    try:
        reportData.validStartDate = datetime.date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        reportData.validEndDate = datetime.date(int(match.group(4)), int(match.group(5)), int(match.group(6)))
    except Exception as e:
        reportData.validStartDate = datetime.datetime.now().date()
        reportData.validEndDate = datetime.datetime.now().date()
        traceback.print_exc()
    # 检测日期
    reportData.checkDate = reportData.validStartDate
    # 项目地址
    reportData.projAddr = delSpace(_mainTable.rows[_proof_level_index].cells[5].text)
    # 检测结论
    _result_index = find_row_by_first_col_text(_mainTable, "检测结论")
    reportData.reportResult = delSpace(_mainTable.rows[_result_index].cells[1].text)
    # 检测数量
    reportData.checkedUnitCount = count_checked_unit(document)

    # 检测人员（多行）
    reportData.checkMans = []
    j = 0
    checkManOut = ["", "姓名"]
    while j < len(_mainTable.rows):
        row = _mainTable.rows[j]
        cell = row.cells[0]
        # 检查单元格中是否包含目标文字
        if "检测人员" in cell.text:
            checkManId = delSpace(row.cells[5].text)
            checkManName = delSpace(row.cells[1].text)
            if checkManName not in checkManOut:
                reportData.checkMans.append(CheckMan(checkManId, checkManName))
        j += 1
    # 点数
    reportData.checkPointCount = count_check_point(document)
    return reportData
